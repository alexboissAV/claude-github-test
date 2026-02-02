#!/usr/bin/env python3
"""Script to create a Word document with a test title."""

try:
    from docx import Document
    from docx.shared import Pt

    # Create a new Document
    doc = Document()

    # Add a title
    title = doc.add_heading('Test Title', level=1)

    # Save the document
    doc.save('test_document.docx')
    print("Word document created successfully!")

except ImportError:
    print("python-docx library not available. Installing...")
    import subprocess
    import sys

    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

    # Try again after installation
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    title = doc.add_heading('Test Title', level=1)
    doc.save('test_document.docx')
    print("Word document created successfully!")
