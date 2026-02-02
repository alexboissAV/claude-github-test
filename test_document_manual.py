#!/usr/bin/env python3
"""
Alternative script using python-docx library (already installed in your workflow).
This creates a professional Word document with a test title.
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Create a new Document
    doc = Document()

    # Add a title with formatting
    title = doc.add_heading('Test Title', level=1)
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    # Add a paragraph
    paragraph = doc.add_paragraph('This is a test document created by Claude.')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Save the document
    doc.save('test_document.docx')
    print("Successfully created test_document.docx with formatting!")

except ImportError:
    print("python-docx not available, falling back to standard library approach")
    # Fall back to the standard library approach
    import sys
    sys.path.insert(0, '/home/runner/work/claude-github-test/claude-github-test')
    from create_word_doc import create_word_document
    create_word_document()
