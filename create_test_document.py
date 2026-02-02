#!/usr/bin/env python3
"""
Script to create a Word document with a test title.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_document():
    # Create a new Document
    doc = Document()

    # Add a title
    title = doc.add_heading('Test Title', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add some content
    doc.add_paragraph('This is a test Word document created by Claude.')
    doc.add_paragraph('Document generated on: 2026-02-02')

    # Save the document
    doc.save('test_document.docx')
    print('✅ Word document created successfully: test_document.docx')

if __name__ == '__main__':
    create_test_document()
